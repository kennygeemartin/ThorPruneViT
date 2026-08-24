"""Import real experiment summaries into a ThorPruneViT DOCX.

This script intentionally refuses to run if the expected real-run summary files are
missing. It never substitutes smoke-test or synthetic results.
"""
import argparse,re,sys
from pathlib import Path
import pandas as pd
from docx import Document
from common import ROOT

ap=argparse.ArgumentParser(); ap.add_argument('--input',required=True); ap.add_argument('--output',required=True); args=ap.parse_args()
rev=ROOT/'results'/'reviewer_tables'; summary=rev/'summary_mean_sd_ci.csv'
if not summary.exists():
    raise SystemExit('Real experiment summaries are missing. Run the full real-data experiments and build_reviewer_tables.py first. Synthetic smoke-test results are never imported.')
s=pd.read_csv(summary)

def stat(exp_contains, metric):
    q=s[s['experiment'].str.contains(exp_contains,regex=False) & (s['metric']==metric)]
    if q.empty: return None
    r=q.iloc[0]; return f"{r['mean']:.4f} ± {r['sd']:.4f}"

doc=Document(args.input)
# Update existing aggregate tables conservatively when matches exist.
# Table 1: baseline training/validation values are left unchanged because the rerun
# scripts currently summarize test metrics. Table 3 efficiency can be updated.
if len(doc.tables)>=3:
    t=doc.tables[2]
    # baseline/pruned parameters and FLOPs
    bpar=stat('baseline_seed_*','parameters'); bfl=stat('baseline_seed_*','approx_flops')
    ppar=stat('target_0.53_tw_0.50_mw_0.50_joint','parameters'); pfl=stat('target_0.53_tw_0.50_mw_0.50_joint','approx_flops')
    blat=stat('baseline_seed_*','latency_ms_mean'); plat=stat('target_0.53_tw_0.50_mw_0.50_joint','latency_ms_mean')
    vals=[bpar,bfl,blat,ppar,pfl,plat]
    if all(v is not None for v in vals):
        t.rows[1].cells[1].text=bpar; t.rows[1].cells[2].text=bfl; t.rows[1].cells[3].text=blat
        t.rows[2].cells[1].text=ppar; t.rows[2].cells[2].text=pfl; t.rows[2].cells[3].text=plat

# Append reviewer-required tables as an explicit Results addendum.
doc.add_heading('Reviewer-required reproducibility results',level=2)
for fname,title in [
    ('summary_mean_sd_ci.csv','Multi-seed aggregate results'),
    ('nih_baseline_disease_wise_summary.csv','NIH disease-wise baseline results'),
    ('nih_thorprunevit_disease_wise_summary.csv','NIH disease-wise ThorPruneViT results')]:
    fp=rev/fname
    if not fp.exists(): continue
    df=pd.read_csv(fp)
    p=doc.add_paragraph(); p.add_run(title).bold=True
    tab=doc.add_table(rows=1,cols=len(df.columns)); tab.style='Table Grid'
    for i,c in enumerate(df.columns): tab.rows[0].cells[i].text=str(c)
    for _,row in df.iterrows():
        cells=tab.add_row().cells
        for i,v in enumerate(row): cells[i].text='' if pd.isna(v) else str(v)

doc.save(args.output); print(args.output)
